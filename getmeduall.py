import streamlit as st
import pandas as pd
import requests
from bs4 import BeautifulSoup
import re
import time
from io import BytesIO
from docx import Document
from docx.shared import Inches

# ⭐ GitHubのCSV URL（orangeちゃんのGitHub URLに差し替えてね）
CSV_URL = "https://raw.githubusercontent.com/orangenyan20/Getmeduall/main/data.csv"

@st.cache_data
def load_topic_mapping():
    try:
        df = pd.read_csv(CSV_URL)
        mapping = dict(zip(df.iloc[:, 0].astype(str), df.iloc[:, 1]))
        return mapping
    except Exception as e:
        st.error(f"単元データの読み込みに失敗しました: {e}")
        return {}

def get_page_text(url, get_images=True):
    try:
        resp = requests.get(url)
        if resp.status_code != 200:
            return None
        soup = BeautifulSoup(resp.text, 'html.parser')
        problem = soup.find('div', class_='quiz-body mb-64')
        choices = [f"{c.find('span', class_='choice-header').text.strip()} {c.find_all('span')[1].text.strip()}"
                   for c in soup.find_all('div', class_='box-select')]
        h4s = soup.find_all('h4')
        ans = h4s[0].text.strip() if h4s else '解答なし'
        qid = re.search(r'([0-9]{3}[A-Za-z][0-9]+)', h4s[1].text).group(1) if len(h4s) >=2 else '問題番号なし'
        expl = soup.find('div', class_='explanation').text.strip() if soup.find('div', class_='explanation') else '解説なし'
        imgs = []

        if get_images:
            for d in soup.find_all('div', class_='box-quiz-image'):
                for a in d.find_all('a', href=True):
                    href = a['href']
                    if '.jpg' in href:
                        imgs.append(href.replace('thumb_', ''))

        return {
            "question_id": qid,
            "problem": problem.text.strip() if problem else '問題文なし',
            "choices": choices,
            "answer": ans,
            "explanation": expl,
            "images": imgs
        }
    except:
        return None

def create_word_doc(pages, year, label, topic_map, include_images=True):
    doc = Document()
    doc.add_heading(f'{year}年 医師国家試験問題（{label}）', 0)
    doc.add_paragraph(f"取得問題数: {len(pages)}問")
    for i, p in enumerate(pages, 1):
        doc.add_heading(f"問題{ i } {p['question_id']}", level=2)
        unit = topic_map.get(p['question_id'], "分野名なし")
        doc.add_paragraph(f"分野: {unit}")
        doc.add_paragraph(p['problem'])
        if include_images and p['images']:
            for url in p['images']:
                try:
                    r = requests.get(url)
                    if r.status_code == 200:
                        img_stream = BytesIO(r.content)
                        doc.add_picture(img_stream, width=Inches(2.5))
                except:
                    pass
        doc.add_paragraph("選択肢：")
        for c in p['choices']:
            doc.add_paragraph(c)
        doc.add_paragraph(p['answer'])
        doc.add_paragraph("解説: " + p['explanation'])
        doc.add_page_break()
    fn = f"{year}_{label}_medu4.docx"
    doc.save(fn)
    return fn

def scrape_sections(year, sections, topic_map, include_images=True):
    collected = []
    for sec in sections:
        st.markdown(f"### ▶️ セクション {sec} を取得中...")
        bar = st.progress(0)
        fail_count = 0
        for i, num in enumerate(range(1, 81)):
            qid = f"{year}{sec}{num}"
            url = f"https://medu4.com/{qid}"
            data = get_page_text(url, get_images=include_images)
            if data:
                data["question_id"] = qid
                collected.append(data)
                fail_count = 0
            else:
                fail_count += 1
                if fail_count >= 3:
                    st.warning(f"⚠ {sec}セクションで3問連続失敗 → スキップします")
                    break
            bar.progress((i + 1) / 80)
            time.sleep(0.15)
        bar.empty()
    return collected

# 🎛️ UI
st.title("🩺 国試問題取得ツール（GitHub連携版）")
year = st.text_input("年度を入力（例: 100）")
include_images = st.checkbox("画像も取得する", value=True)

# GitHubから単元データ取得
topic_map = load_topic_mapping()

col1, col2 = st.columns(2)

with col1:
    if st.button("A〜Cセクション取得開始"):
        if year:
            ac_data = scrape_sections(year, ['A', 'B', 'C'], topic_map, include_images)
            if ac_data:
                fn_ac = create_word_doc(ac_data, year, "A-C", topic_map, include_images)
                st.success("✅ A〜Cセクション完了！")
                with open(fn_ac, "rb") as f:
                    st.download_button("A〜CのWordファイルをダウンロード", f, file_name=fn_ac)
            else:
                st.error("❌ A〜Cセクションで有効な問題が取得できませんでした。")

with col2:
    if st.button("D〜Iセクション取得開始"):
        if year:
            di_data = scrape_sections(year, ['D', 'E', 'F', 'G', 'H', 'I'], topic_map, include_images)
            if di_data:
                fn_di = create_word_doc(di_data, year, "D-I", topic_map, include_images)
                st.success("✅ D〜Iセクション完了！")
                with open(fn_di, "rb") as f:
                    st.download_button("D〜IのWordファイルをダウンロード", f, file_name=fn_di)
            else:
                st.error("❌ D〜Iセクションで有効な問題が取得できませんでした。")
